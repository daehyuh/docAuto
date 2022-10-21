import docx
import glob
from docx.shared import Cm, Inches, Mm

folder = "input/6"

img_list = glob.glob(folder+'/*.png')

list(set(img_list))
img_list.sort()
print(img_list, len(img_list))

document = docx.Document()

tbl = document.add_table(rows=0, cols=1)
tbl.style = document.styles['Table Grid']

for i in img_list:
    print(i)
    row_cells = tbl.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(i, width= Mm(54.57), height= Mm(30.93))

document.save("C:/Users/daehy/OneDrive/바탕 화면/새 폴더/7,8/장애물 결과.docx")
print('end')
